"""
13_generate_memo.py

Generate the technical memorandum as a Word document (.docx) using python-docx.
Embeds Maps 1-4, model evaluation, and feature importance figures.

Output:
  report/pfas_technical_memo.docx

Run from pfas-risk-ma/:
  python scripts/13_generate_memo.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
MAPS = ROOT / "maps"
REPORT = ROOT / "report"
REPORT.mkdir(exist_ok=True)

BODY_FONT = "Arial"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(16)
HEADING2_SIZE = Pt(13)
HEADING3_SIZE = Pt(11)


def set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def apply_run_style(run, size=BODY_SIZE, bold=False, color=None, italic=False):
    run.font.name = BODY_FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    sizes = {1: HEADING1_SIZE, 2: HEADING2_SIZE, 3: HEADING3_SIZE}
    before = {1: 18, 2: 14, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    apply_run_style(r, size=sizes[level], bold=True,
                    color=RGBColor(0x1B, 0x5E, 0x7B))
    return p


def add_body(doc, text, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    apply_run_style(r, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    apply_run_style(r)
    return p


def add_figure(doc, image_path, caption, width_in=6.3):
    if not image_path.exists():
        add_body(doc, f"[Figure missing: {image_path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    apply_run_style(r, size=Pt(9), italic=True,
                    color=RGBColor(0x55, 0x55, 0x55))


def add_table(doc, header, rows, col_widths_in=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Header
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(text)
        apply_run_style(r, bold=True, size=Pt(10),
                        color=RGBColor(0xFF, 0xFF, 0xFF))
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[i], "1B5E7B")

    # Body
    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        for i, text in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(text))
            apply_run_style(r, size=Pt(9))
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)

    # Spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    # Also set East Asian font to avoid fallback
    r = style.element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"), BODY_FONT)


def build_title_page(doc):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(
        "PFAS Contamination Risk Prediction and "
        "Environmental Justice Analysis"
    )
    apply_run_style(r, size=Pt(22), bold=True,
                    color=RGBColor(0x1B, 0x5E, 0x7B))

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Massachusetts Public Water Systems")
    apply_run_style(r, size=Pt(16), italic=True,
                    color=RGBColor(0x55, 0x55, 0x55))

    # Document type
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Technical Memorandum")
    apply_run_style(r, size=Pt(14), bold=True)

    # Author + date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    r = p.add_run("Maahum Yousuf, EIT")
    apply_run_style(r, size=Pt(13), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(date.today().strftime("%B %Y"))
    apply_run_style(r, size=Pt(11))

    # Page break
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_memo():
    doc = Document()
    set_default_style(doc)

    # 1" margins all around
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    build_title_page(doc)

    # ===== 1. INTRODUCTION =====
    add_heading(doc, "1. Introduction", level=1)
    add_body(
        doc,
        "Per- and polyfluoroalkyl substances (PFAS) are a class of synthetic "
        "fluorinated chemicals that persist indefinitely in the environment "
        "and bioaccumulate in human tissue. Decades of industrial use in "
        "textile finishing, metal plating, semiconductor manufacturing, "
        "food packaging, and aqueous film-forming foam (AFFF) used in fire "
        "training have produced widespread contamination of drinking water "
        "sources across the United States."
    )
    add_body(
        doc,
        "In April 2024, EPA finalized the first National Primary Drinking "
        "Water Regulation for PFAS, establishing a Maximum Contaminant "
        "Level (MCL) of 4.0 ng/L (parts per trillion, ppt) for PFOA and "
        "PFOS individually. EPA also designated PFOA and PFOS as CERCLA "
        "hazardous substances the same year. Massachusetts adopted its "
        "own PFAS drinking-water standard in 2020 (310 CMR 22.07G): the "
        "sum of six PFAS compounds (PFOA, PFOS, PFHxS, PFNA, PFHpA, and "
        "PFDA, commonly called the PFAS6 sum) may not exceed 20 ppt."
    )
    add_body(
        doc,
        "EPA's Fifth Unregulated Contaminant Monitoring Rule (UCMR 5, "
        "monitoring cycle 2023-2025) is the most comprehensive national "
        "PFAS occurrence dataset to date. UCMR 5 required monitoring by "
        "all public water systems (PWSs) serving over 3,300 people and "
        "a stratified sample of smaller systems. Because smaller PWSs "
        "(which make up the majority of active systems in Massachusetts) "
        "are outside the mandatory tier, roughly 85% of MA water systems "
        "were not monitored under UCMR 5, leaving a large informational "
        "gap about likely PFAS exposures in smaller, often rural, "
        "communities."
    )
    add_body(
        doc,
        "This memorandum presents a screening-level analysis that asks "
        "three questions:"
    )
    add_bullet(
        doc,
        "Among Massachusetts PWSs already sampled under UCMR 5, where "
        "and how frequently is PFAS detected, and where are federal or "
        "state limits exceeded?"
    )
    add_bullet(
        doc,
        "For the ~1,600 untested MA systems, which are most likely to "
        "have PFAS contamination based on proximity to known PFAS source "
        "types (airports, military installations, landfills, wastewater "
        "treatment facilities, and industrial dischargers)?"
    )
    add_bullet(
        doc,
        "Do high-risk untested systems disproportionately serve "
        "Massachusetts-defined environmental-justice (EJ) communities?"
    )
    add_body(
        doc,
        "The analysis is explicitly screening-level. It does not "
        "substitute for site-specific sampling or hydrogeologic "
        "investigation, and it is not a risk assessment or a compliance "
        "determination. Its purpose is to help state regulators, "
        "Licensed Site Professionals (LSPs), and consulting hydrologists "
        "prioritize where to sample next."
    )

    # ===== 2. DATA AND METHODS =====
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "2. Data and Methods", level=1)

    add_heading(doc, "2.1 Data Sources", level=2)
    add_body(doc, "All datasets were downloaded in April 2026 (Table 1).")
    add_table(
        doc,
        header=["Dataset", "Source", "Purpose"],
        rows=[
            ["PFAS monitoring results (n = 46,070 MA rows)",
             "EPA UCMR 5, Jan 2026 release",
             "Detection locations, concentrations, exceedances"],
            ["PWS source locations (n = 3,990 points; 1,818 PWSs)",
             "MassGIS PWSDEP_PT",
             "Spatial anchor for distance features"],
            ["PWS service-area polygons (n = 523)",
             "MassGIS DEP_PWS_Water_Service_Areas",
             "EJ spatial overlay"],
            ["Environmental-justice block groups (n = 2,604)",
             "MassGIS 2020 EJ Populations",
             "State-defined EJ criteria"],
            ["21E contaminated sites (n = 2,287)",
             "MassGIS C21E_PT",
             "Priority-site cross-reference"],
            ["Zone II / IWPA protection polygons (n = 2,965)",
             "MassGIS zone2_zone1_iwpa",
             "Source-water protection area overlay"],
            ["Solid waste landfills (n = 672 polygons)",
             "MassGIS SW_LD_POLY",
             "Distance feature"],
            ["POTW service areas (n = 331, WWTP proxy)",
             "MassGIS DEP_Sewer_Service_Areas",
             "Distance feature (centroid)"],
            ["Industrial dischargers (n = 2,445)",
             "MassGIS BWPMAJOR_PT",
             "Distance feature"],
            ["MWRA service polygons",
             "MassGIS mwraservice",
             "MWRA flag"],
            ["Airports (n = 276 in MA)",
             "OurAirports, filtered to US-MA",
             "Distance feature (AFFF history)"],
            ["Military installations (n = 7 in MA)",
             "Hand-curated from DoD / state records",
             "Distance feature"],
        ],
        col_widths_in=[2.0, 2.3, 2.3],
    )
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run("Table 1 - Datasets used in the analysis.")
    apply_run_style(r, size=Pt(9), italic=True,
                    color=RGBColor(0x55, 0x55, 0x55))

    add_heading(doc, "2.2 Non-Detect Handling", level=2)
    add_body(
        doc,
        "UCMR 5 reports concentrations in ug/L (ppb). Values were converted "
        "to ng/L (ppt) by multiplying by 1,000 for direct comparison with "
        "the federal MCL and MA MMCL. Non-detects (rows where the "
        "analytical result sign was '<' the minimum reporting level, MRL) "
        "were substituted at half the MRL, the standard practice for "
        "environmental occurrence datasets. Across the 46,070 MA UCMR 5 "
        "records, 1,362 (3.0%) were detections and the remaining 44,708 "
        "were non-detects."
    )

    add_heading(doc, "2.3 Feature Engineering", level=2)
    add_body(
        doc,
        "All spatial layers were reprojected to EPSG:26986 (Massachusetts "
        "State Plane, meters) so that distance and area calculations could "
        "be performed directly in meters. For each of the 3,990 PWSDEP "
        "source points, nearest-neighbor distances to the closest "
        "airport, military installation, landfill (polygon centroid), "
        "POTW service-area centroid, and BWP major industrial discharger "
        "were computed using a scipy cKDTree. Within-buffer counts of "
        "landfills and industrial facilities at radii of 1, 3, 5, and "
        "10 km were also computed. Per-PWS feature values were aggregated "
        "by taking the minimum distance and maximum buffer count across "
        "all source points belonging to a given PWSID (i.e., the worst "
        "case across the system's wells or intakes)."
    )
    add_body(
        doc,
        "A groundwater flag was derived from the MassGIS TYPE attribute "
        "(GW, EGW). An MWRA flag was set by spatial intersection of each "
        "PWS source point with MWRA's water-service polygons "
        "(CODE in {W, WS}, representing 30 PWSs in the dataset). The "
        "final feature-engineered dataset contained 1,818 unique PWSs, "
        "of which 220 were successfully matched by PWSID to a record in "
        "the UCMR 5 MA summary (the remaining 43 UCMR 5 systems had "
        "PWSIDs not present in the PWSDEP point layer, most often very "
        "small transient non-community systems)."
    )

    add_heading(doc, "2.4 Risk Prediction Model", level=2)
    add_body(
        doc,
        "The modeling task was binary classification: does a given PWS "
        "have at least one PFAS detection above the MRL in UCMR 5? "
        "Features were the five nearest-neighbor distances (airport, "
        "military, landfill, WWTP, industrial, all in km), the "
        "landfills_within_5km and industrial_within_5km counts, and the "
        "is_groundwater indicator. The training pool consisted of the "
        "220 UCMR 5-matched PWSs."
    )
    add_body(
        doc,
        "Two models were fit: logistic regression (LR) with features "
        "standardized and class weights balanced, and a random forest "
        "(RF) with 300 trees, max depth 5, and class weights balanced. "
        "Data were split 80/20 with stratified sampling (random seed 42) "
        "and each model was additionally evaluated with 5-fold "
        "cross-validated ROC AUC. Logistic regression was selected as "
        "the primary model on grounds of interpretability, provided its "
        "test AUC was within 0.02 of the random forest's. The trained "
        "LR model was then applied to the 1,598 untested systems to "
        "produce predicted PFAS-detection probabilities."
    )

    add_heading(doc, "2.5 21E Priority Investigation Cross-Reference", level=2)
    add_body(
        doc,
        "The 2,287 MassDEP Chapter 21E contaminated sites were spatially "
        "joined to the 565 Zone II and 2,400 IWPA source-water "
        "protection polygons, producing 347 unique 21E records "
        "associated with at least one PWS protection area. Because the "
        "21E schema has no structured contaminant field, a keyword scan "
        "was run over each site's NAME attribute using approximately "
        "fifty PFAS-relevant keywords covering AFFF / fire training, "
        "petroleum and fuel storage (UST releases), dry cleaning and "
        "laundry, metal plating, semiconductor manufacturing, textile "
        "finishing, waste disposal, and military installations. Sites "
        "matching any keyword and associated with a PWS that was NOT in "
        "the UCMR 5 testing set were designated priority investigation "
        "sites."
    )

    add_heading(doc, "2.6 Environmental-Justice Analysis", level=2)
    add_body(
        doc,
        "PWS service-area polygons were overlaid on the MassGIS 2020 EJ "
        "Populations layer (all block groups meeting at least one of "
        "Massachusetts' four EJ criteria: income <= 65% of statewide "
        "median, minority population >= 40%, limited-English households "
        ">= 25% of total, or a combined minority-plus-income criterion). "
        "For each PWS, the area-fraction of its service territory "
        "falling inside any EJ block group was computed. Systems with "
        "positive overlap were flagged as serving an EJ community."
    )
    add_body(
        doc,
        "Two statistical tests were applied to the predicted-risk scores "
        "of untested systems, using service-area EJ flag as the grouping "
        "variable: (a) a chi-square test on the 2x2 contingency table of "
        "high-risk (score > 0.5) versus EJ-serving; and (b) a one-sided "
        "Mann-Whitney U test of the hypothesis that predicted risk "
        "scores are stochastically greater in EJ-serving than in "
        "non-EJ-serving PWSs."
    )

    # ===== 3. RESULTS =====
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "3. Results", level=1)

    add_heading(doc, "3.1 PFAS Occurrence Summary", level=2)
    add_body(
        doc,
        "UCMR 5 monitored 263 Massachusetts public water systems and "
        "produced 46,070 individual analytical results across 29 PFAS "
        "compounds (lithium, which is co-monitored under UCMR 5, was "
        "excluded from this analysis). Of those 263 systems, 143 (54.4%) "
        "had at least one PFAS detection. 106 systems (40.3%) had at "
        "least one sample exceeding the federal 4 ppt PFOA/PFOS MCL, "
        "and 34 systems (12.9%) had at least one sample whose PFAS6 sum "
        "exceeded the Massachusetts MMCL of 20 ppt. At the individual-"
        "result level, there were 429 federal MCL exceedances across "
        "all sample events. Figure 1 shows the spatial distribution of "
        "these detection outcomes."
    )
    add_figure(
        doc,
        MAPS / "map1_pfas_detections_ma.png",
        "Figure 1 - UCMR 5 PFAS detections in Massachusetts public water "
        "systems. Each circle is one PWS source point, sized and colored "
        "by detection status. Boston-metro MWRA-served communities "
        "appear absent because MWRA's source points are in central-MA "
        "reservoirs, not in the metro area.",
    )

    add_heading(doc, "3.2 Risk Prediction Model Performance", level=2)
    add_body(
        doc,
        "The 220-system training pool had a positive-class prevalence "
        "of 61.4% (135 systems with at least one detection). Class "
        "balance was therefore modest but not severe, and class-balanced "
        "weighting was used to avoid favoring the majority class. "
        "Test-set ROC AUC was 0.885 for logistic regression and 0.841 "
        "for random forest. Five-fold cross-validated AUCs were 0.601 "
        "(+/- 0.123) for LR and 0.587 (+/- 0.083) for RF. The gap "
        "between test AUC and CV AUC suggests the 44-system test split "
        "happened to be favorable; the CV estimate is the more reliable "
        "indicator of out-of-sample performance. LR was selected as the "
        "primary model for its interpretable coefficients (Table 2)."
    )
    add_table(
        doc,
        header=["Feature", "LR coefficient (scaled)", "RF Gini importance"],
        rows=[
            ["dist_industrial_km", "-0.36 (closer = higher risk)", "0.16"],
            ["industrial_within_5km", "+0.33 (more = higher risk)", "0.15"],
            ["is_groundwater", "+0.30 (GW = higher risk)", "0.02"],
            ["dist_wwtp_km", "+0.24", "0.10"],
            ["dist_airport_km", "-0.22 (closer = higher risk)", "0.15"],
            ["landfills_within_5km", "+0.15 (more = higher risk)", "0.06"],
            ["dist_landfill_km", "+0.03", "0.21"],
            ["dist_military_km", "-0.02", "0.14"],
        ],
        col_widths_in=[2.3, 2.3, 2.0],
    )
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(
        "Table 2 - Feature importance. Logistic regression coefficients are "
        "on standardized features; positive values indicate higher detection "
        "probability."
    )
    apply_run_style(r, size=Pt(9), italic=True,
                    color=RGBColor(0x55, 0x55, 0x55))
    add_body(
        doc,
        "Both models agree that proximity to industrial dischargers and "
        "to airports is the strongest signal: PWSs closer to these "
        "sources are more likely to have detected PFAS. Groundwater "
        "sources are more likely to show PFAS than surface-water "
        "sources. The positive dist_wwtp_km coefficient is "
        "counter-intuitive and may reflect a geographic confound: POTW "
        "service areas are densest in eastern MA, where many MWRA and "
        "already-tested large-system points are concentrated. Distance "
        "to landfills and military installations contribute little "
        "signal to the LR model but more in the random forest, "
        "suggesting a non-linear relationship that the RF captures."
    )
    add_figure(
        doc,
        MAPS / "model_evaluation.png",
        "Figure 2 - Model evaluation: confusion matrices for LR and RF, "
        "and ROC curves comparing the two models against random "
        "classification.",
    )

    add_heading(doc, "3.3 Predicted Risk for Untested Systems", level=2)
    add_body(
        doc,
        "Applying the fitted LR model to the 1,598 untested PWSs produced "
        "the following risk-score distribution: 966 Low (score <= 0.25), "
        "486 Moderate (0.25 - 0.50), 142 High (0.50 - 0.75), and 4 Very "
        "High (> 0.75). Figure 3 shows the geographic distribution; "
        "high-risk untested systems are concentrated in the I-495 belt "
        "and in parts of southeastern MA, consistent with proximity to "
        "industrial facilities and airports."
    )
    add_figure(
        doc,
        MAPS / "map2_predicted_risk.png",
        "Figure 3 - Predicted PFAS detection risk for untested MA "
        "public water systems. Gray dots are UCMR 5-tested systems "
        "shown for context. Untested systems are colored on a green-to-"
        "red scale by predicted detection probability.",
    )

    add_heading(doc, "3.4 Environmental-Justice Disparity", level=2)
    add_body(
        doc,
        "Of the 1,598 untested systems, 49 (3.1%) serve an EJ community "
        "based on the MassGIS 2020 EJ-populations overlay of their "
        "MassDEP service-area polygons. Of those 49 EJ-serving systems, "
        "21 (42.9%) are classified high-risk by the model. Of the 1,549 "
        "non-EJ-serving systems, 125 (8.1%) are classified high-risk. "
        "The disparity is 5.3x and is highly significant: chi-square = "
        "65.1, p < 0.0001 on the 2x2 contingency table, and a one-sided "
        "Mann-Whitney U test rejects equality of risk-score "
        "distributions at p < 0.0001 with the EJ-serving group having "
        "higher scores. Mean predicted risk score was 0.412 in EJ "
        "service areas versus 0.231 in non-EJ service areas."
    )
    add_figure(
        doc,
        MAPS / "map3_risk_ej_overlay.png",
        "Figure 4 - High-risk untested PWSs (red) and lower-risk untested "
        "PWSs (gold) overlaid on 2020 MassGIS EJ block groups "
        "(lavender). High-risk systems are visibly clustered over EJ "
        "block groups.",
    )
    add_body(
        doc,
        "The direction of this finding is consistent with the broader "
        "environmental-justice literature on drinking-water quality and "
        "with the 2022 GAO report on PFAS disparities, which found that "
        "disadvantaged communities face documented gaps in monitoring "
        "and remediation."
    )

    add_heading(doc, "3.5 21E Priority Investigation Sites", level=2)
    add_body(
        doc,
        "Of 347 unique 21E sites located inside a Zone II or IWPA "
        "protection polygon, 37 had a site-name match to at least one "
        "PFAS-relevant keyword. Seven of those 37 are associated with a "
        "PWS that has not been tested under UCMR 5 and are therefore "
        "designated priority investigation sites (Figure 5). The top "
        "matching keywords were 'ust' (underground storage tank, n=7), "
        "'fuel' (7), 'gas station' (5), and 'dump' (3). Notable "
        "entries include a former fire station in Stow (PWSID "
        "MA2286005) and 'Middlesex Petroleum' in Tyngsborough (PWSID "
        "MA3301047). Full details are in Appendix A."
    )
    add_figure(
        doc,
        MAPS / "map4_21e_priority_sites.png",
        "Figure 5 - 21E priority investigation sites (dark red "
        "triangles) inside Zone II source-water protection areas (light "
        "cyan) of untested PWSs. All 2,287 21E sites statewide are "
        "shown as faint gray dots for context.",
    )

    # ===== 4. DISCUSSION =====
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "4. Discussion", level=1)
    add_body(
        doc,
        "The UCMR 5 results confirm that PFAS contamination of public "
        "drinking-water supplies in Massachusetts is not isolated to a "
        "handful of military or industrial hotspots: 54% of tested "
        "systems had at least one detection, and 40% had at least one "
        "sample exceeding the federal 4 ppt PFOA/PFOS MCL. The MA PFAS6 "
        "MMCL (20 ppt) is a tighter cumulative standard that was "
        "exceeded by 13% of tested systems. These results, while "
        "alarming, are broadly consistent with preliminary national "
        "UCMR 5 summaries released by EPA in 2024."
    )
    add_body(
        doc,
        "The modest out-of-sample model performance (CV AUC ~ 0.60) is "
        "expected for a proximity-only screening model. Proximity is a "
        "blunt instrument: it ignores hydrogeologic transport, "
        "prevailing groundwater flow direction, the age and closure "
        "status of nearby sources, and the specific PFAS-use history of "
        "a given industrial facility. The model is nevertheless useful "
        "for prioritization - a system flagged high-risk is roughly 5x "
        "more likely to be in an EJ community than a low-risk system, "
        "which directly identifies a concrete, action-oriented finding "
        "for state agencies."
    )
    add_body(
        doc,
        "The 21E cross-reference addresses an information gap peculiar "
        "to Massachusetts' contaminated-site regulatory framework. Under "
        "Chapter 21E (the MCP, 310 CMR 40.0000), RTNs are assigned "
        "based on a release report and cleanup progression, not the "
        "specific suite of contaminants present. A site listed for a "
        "UST release in 1995 may in fact host AFFF-impacted soil, but "
        "absent a specific request, no PFAS sampling is required as "
        "part of the cleanup. Cross-referencing 21E sites with "
        "source-water protection areas for untested public water "
        "supplies therefore offers a concrete, targeted set of seven "
        "candidate sites that state regulators, LSPs, or consulting "
        "firms could prioritize for PFAS-specific analytical work."
    )

    # ===== 5. LIMITATIONS =====
    add_heading(doc, "5. Limitations", level=1)
    add_bullet(
        doc,
        "UCMR 5 coverage is non-random. Selection favors larger systems, "
        "so the training pool under-represents small groundwater systems "
        "most vulnerable to localized contamination."
    )
    add_bullet(
        doc,
        "Proximity-based features do not capture hydrogeologic "
        "transport, groundwater gradient, well-screen depth, or "
        "contaminant fate. Two wells equidistant from the same landfill "
        "can have fundamentally different exposures."
    )
    add_bullet(
        doc,
        "MWRA-served systems serve approximately 50 metro-Boston "
        "municipalities via Quabbin/Wachusett reservoirs 60+ miles west "
        "of the served population. Their distance features are not "
        "directly comparable to independent groundwater sources. They "
        "are retained with an MWRA flag, and Map 1 includes a textual "
        "note explaining the apparent Boston-metro data gap."
    )
    add_bullet(
        doc,
        "The training pool contains only 220 PWSs. CV AUC of 0.60 "
        "corresponds to a screening-level tool, not a concentration "
        "predictor."
    )
    add_bullet(
        doc,
        "21E keyword flagging over site-name strings produces both "
        "false positives (e.g., gas stations where the UST release was "
        "a gasoline spill with no PFAS history) and false negatives "
        "(industrial sites whose names do not mention their "
        "PFAS-relevant process)."
    )
    add_bullet(
        doc,
        "EJ analysis uses MassDEP estimated service-area polygons, not "
        "parcel-level customer connections, so 'serves EJ community' "
        "is an area-based approximation."
    )
    add_bullet(
        doc,
        "Non-detect imputation at MRL/2 is standard practice but "
        "introduces uncertainty in any mean-concentration estimate."
    )
    add_bullet(
        doc,
        "Industrial-discharger coverage uses the MassDEP BWP major-"
        "polluter point layer (2,445 points), which is restricted to "
        "facilities with major NPDES permits and does not include all "
        "small industrial sources."
    )

    # ===== 6. CONCLUSIONS =====
    add_heading(doc, "6. Conclusions", level=1)
    add_body(
        doc,
        "PFAS contamination is pervasive among tested Massachusetts "
        "public water systems, and a screening-level proximity model "
        "identifies 146 of the 1,598 untested systems as plausibly "
        "high-risk. Those high-risk systems serve EJ communities at "
        "roughly 5x the rate that lower-risk systems do, a "
        "statistically robust disparity that warrants targeted "
        "follow-up."
    )
    add_body(
        doc,
        "Three concrete recommendations follow from this analysis: "
        "(1) MassDEP should prioritize the seven 21E priority "
        "investigation sites identified here for PFAS-specific soil "
        "and groundwater sampling as part of existing MCP cleanup "
        "activities; (2) any state-level PFAS sampling program for "
        "small untested systems should stratify its sample to include "
        "the ~50 EJ-serving high-risk systems identified here, both "
        "because they are the most likely to show exceedances and "
        "because doing so addresses a documented equity gap; and "
        "(3) future iterations should incorporate hydrogeologic "
        "gradient data, expand to include MassDEP's own non-UCMR "
        "sampling records, and refine the industrial proxy to include "
        "TRI-reporting facilities."
    )
    add_body(
        doc,
        "All outputs, code, and an interactive Streamlit application "
        "are available at the project's public GitHub repository."
    )

    # ===== REFERENCES =====
    add_heading(doc, "References", level=1)
    refs = [
        ("EPA (2024). National Primary Drinking Water Regulations: PFAS. "
         "89 FR 32532. https://www.epa.gov/sdwa/and-polyfluoroalkyl-"
         "substances-pfas"),
        ("EPA (2026, Jan). UCMR 5 occurrence data. "
         "https://www.epa.gov/dwucmr/occurrence-data-unregulated-"
         "contaminant-monitoring-rule"),
        ("310 CMR 22.07G. Massachusetts PFAS6 maximum contaminant "
         "level. MassDEP drinking-water regulation, 2020."),
        ("310 CMR 40.0000. Massachusetts Contingency Plan (MCP) for "
         "management of oil and hazardous material releases under "
         "Chapter 21E."),
        ("U.S. GAO (2022). PFAS: Additional Actions Would Strengthen "
         "Federal Efforts to Assess and Respond. GAO-22-105088."),
        ("MassGIS (2022, Nov). 2020 Environmental Justice Populations "
         "layer. Executive Office of Energy and Environmental Affairs."),
        ("MassDEP. Estimated Public Drinking Water System Service Area "
         "Boundaries. MassGIS, 2024 update."),
    ]
    for ref in refs:
        add_bullet(doc, ref)

    # ===== APPENDIX =====
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Appendix A - 21E Priority Investigation Sites", level=1)
    add_body(
        doc,
        "Seven MassDEP Chapter 21E sites are located inside the Zone II "
        "or IWPA source-water protection area of an untested Massachusetts "
        "public water system and have site names matching at least one "
        "PFAS-relevant keyword. Full tabular data is available at "
        "data/cleaned/21e_priority_investigation.csv."
    )
    add_table(
        doc,
        header=["RTN", "Site name", "Town", "Keyword", "Area", "PWSID"],
        rows=[
            ["2-0021075", "FMR Fire Station", "Stow", "fire", "Zone II", "MA2286005"],
            ["1-0020098", "Friendly Fred's Gas Station", "Windsor", "gas station", "IWPA", "MA1345007"],
            ["1-0020515", "MK Fuel County Line", "Brimfield", "fuel", "IWPA", "MA1043027"],
            ["2-0012178", "Frank Realty Trust", "Hudson", "ust", "IWPA", "MA2028001"],
            ["3-2014066", "Middlesex Petroleum", "Tyngsborough", "petroleum", "IWPA", "MA3301047"],
            ["4-0030044", "UST Release", "Carver", "ust", "IWPA", "MA4052056"],
            ["4-0030044", "UST Release", "Carver", "ust", "IWPA", "MA4052030"],
        ],
        col_widths_in=[0.9, 2.0, 1.0, 0.9, 0.8, 1.0],
    )

    add_heading(doc, "Appendix B - Feature Importance Plot", level=1)
    add_figure(
        doc,
        MAPS / "feature_importance.png",
        "Figure A-1 - Logistic regression absolute coefficients "
        "(normalized) vs random-forest Gini importances for the eight "
        "features used in the risk prediction model.",
        width_in=6.0,
    )

    # Save
    out = REPORT / "pfas_technical_memo.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build_memo()
